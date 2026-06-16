"""
Word文档导出工具
将优化后的简历、面试问题、评分报告导出为格式化的Word文档(.docx)
"""
import io
import re
from pathlib import Path
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from models.schema import (
    OptimizedResume,
    ScoreReport,
    InterviewQuestion,
    OptimizationLogic,
)
from app_config import config


class WordExporter:
    """
    Word格式导出器

    支持导出：
    1. 优化后的简历
    2. 评分报告
    3. 面试问题列表
    4. 修改日志（Optimization_Logic）
    """

    def __init__(self, output_dir: Path = None):
        """
        初始化导出器

        Args:
            output_dir: 输出目录，默认使用config.OUTPUTS_DIR
        """
        self.output_dir = output_dir or config.OUTPUTS_DIR
        config.ensure_dirs()

    def export_resume(self, resume: OptimizedResume, filename: str = None) -> Path:
        """
        导出优化后的简历为Word

        Args:
            resume: 优化后的简历
            filename: 文件名，默认使用时间戳

        Returns:
            导出文件的路径
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"optimized_resume_{timestamp}.docx"

        filepath = self.output_dir / filename

        doc = Document()
        self._apply_default_styles(doc)
        self._add_resume_content(doc, resume)

        doc.save(str(filepath))
        return filepath

    def export_full_report(
        self,
        resume: OptimizedResume,
        report: ScoreReport,
        questions: List[InterviewQuestion],
        company_name: str = None,
    ) -> Path:
        """
        导出完整报告为Word

        Args:
            resume: 优化后的简历
            report: 评分报告
            questions: 面试问题列表
            company_name: 公司名称

        Returns:
            导出文件的路径
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        company_str = f"_{self._safe_filename_part(company_name)}" if company_name else ""
        filename = f"full_report{company_str}_{timestamp}.docx"

        filepath = self.output_dir / filename

        doc = Document()
        self._apply_default_styles(doc)

        # 标题
        title = doc.add_heading('简历优化完整报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 生成时间
        time_para = doc.add_paragraph()
        time_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # 1. 优化后简历
        doc.add_heading('优化后的简历', level=1)
        self._add_resume_content(doc, resume)
        doc.add_paragraph()

        # 2. 修改说明
        doc.add_heading('修改说明', level=1)
        self._add_optimization_log(doc, resume.optimization_logics)
        doc.add_paragraph()

        # 3. 评分报告
        doc.add_heading('评分报告', level=1)
        self._add_score_report(doc, report)
        doc.add_paragraph()

        # 4. 面试问题
        doc.add_heading('面试问题', level=1)
        self._add_questions(doc, questions)

        doc.save(str(filepath))
        return filepath

    def export_to_bytes(self, resume: OptimizedResume) -> bytes:
        """
        导出简历为字节流（用于API响应）

        Args:
            resume: 优化后的简历

        Returns:
            Word文档的字节流
        """
        doc = Document()
        self._apply_default_styles(doc)
        self._add_resume_content(doc, resume)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def export_full_report_to_bytes(
        self,
        resume: OptimizedResume,
        report: ScoreReport,
        questions: List[InterviewQuestion],
    ) -> bytes:
        """
        导出完整报告为字节流

        Args:
            resume: 优化后的简历
            report: 评分报告
            questions: 面试问题列表

        Returns:
            Word文档的字节流
        """
        doc = Document()
        self._apply_default_styles(doc)

        # 标题
        title = doc.add_heading('简历优化完整报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        time_para = doc.add_paragraph()
        time_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        doc.add_heading('优化后的简历', level=1)
        self._add_resume_content(doc, resume)
        doc.add_paragraph()

        doc.add_heading('修改说明', level=1)
        self._add_optimization_log(doc, resume.optimization_logics)
        doc.add_paragraph()

        doc.add_heading('评分报告', level=1)
        self._add_score_report(doc, report)
        doc.add_paragraph()

        doc.add_heading('面试问题', level=1)
        self._add_questions(doc, questions)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _apply_default_styles(self, doc: Document) -> None:
        """应用默认样式"""
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)

    def _add_resume_content(self, doc: Document, resume: OptimizedResume) -> None:
        """添加简历内容"""
        for section_name, content in resume.sections.items():
            # 章节标题
            heading = doc.add_heading(section_name, level=2)

            # 章节内容
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    p = doc.add_paragraph(line)
                    p.style = 'Normal'

            doc.add_paragraph()

        # 匹配度
        match_para = doc.add_paragraph()
        run = match_para.add_run(f"匹配度评分: {resume.match_score}%")
        run.bold = True
        run.font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)  # 金色

    def _add_optimization_log(self, doc: Document, logics: List[OptimizationLogic]) -> None:
        """添加修改日志"""
        if not logics:
            doc.add_paragraph('（暂无详细修改记录）')
            return

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # 表头
        header_cells = table.rows[0].cells
        headers = ['原文', '修改后', '修改原因', '匹配关键词']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for run in header_cells[i].paragraphs[0].runs:
                run.bold = True

        # 数据行
        for logic in logics:
            row_cells = table.add_row().cells
            row_cells[0].text = logic.original_text
            row_cells[1].text = logic.optimized_text
            row_cells[2].text = logic.reason
            row_cells[3].text = logic.jd_keyword_matched

    def _add_score_report(self, doc: Document, report: ScoreReport) -> None:
        """添加评分报告"""
        # 综合评分
        score_para = doc.add_heading(f"综合评分: {report.overall_score}", level=2)

        readiness_para = doc.add_paragraph()
        run = readiness_para.add_run(f"面试准备度: {report.interview_readiness}")
        run.bold = True

        doc.add_paragraph()

        # 各维度评分
        doc.add_heading('各维度评分', level=3)
        for dimension, score in report.dimensions.items():
            bar = '█' * int(score / 5)
            p = doc.add_paragraph(f'{dimension}: {bar} {score}')

        doc.add_paragraph()

        # 改进建议
        doc.add_heading('改进建议', level=3)
        doc.add_paragraph(report.feedback)

    def _add_questions(self, doc: Document, questions: List[InterviewQuestion]) -> None:
        """添加面试问题"""
        doc.add_paragraph(f'共 {len(questions)} 道面试问题')
        doc.add_paragraph()

        for i, q in enumerate(questions, 1):
            # 问题标题
            doc.add_heading(f'{i}. [{q.type}] {q.question}', level=3)

            # 详情
            p1 = doc.add_paragraph()
            p1.add_run('针对关键词: ').bold = True
            p1.add_run(q.jd_keyword_targeted)

            p2 = doc.add_paragraph()
            p2.add_run('评估标准: ').bold = True
            p2.add_run(q.evaluation_criteria)

            p3 = doc.add_paragraph()
            p3.add_run('参考答案: ').bold = True
            p3.add_run(q.sample_answer)

            doc.add_paragraph()

    def _safe_filename_part(self, value: str) -> str:
        safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value.strip())
        return safe[:40] or "company"
